# import streamlit as st
# import re
# import PyPDF2
# from docx import Document
# from rapidfuzz import fuzz
# st.title("📄 Resume analyzer.py")
# st.markdown("### 👨‍💻 Built by: Piyush Solanki")
# st.markdown("GitHub: https://github.com/piyush-tech243")
# st.markdown("🔗 LinkedIn: https://www.linkedin.com/in/piyush-solanki-9a5bb2317")
# st.markdown("Made with ❤️ using Python, Streamlit, RapidFuzz, PyPDF2 & NLP")

# # resume pdf

# uploaded_file = st.file_uploader("📤 Upload Resume", type=["pdf", "docx", "txt"])
# user_resume = st.text_area("Or paste your Text  here")


# # cleaning data

# def clean_text(text):
#     text = text.lower()               # lowercase
#     text = re.sub(r'\s+', ' ', text)   #remove extra space
#     text = re.sub(r'[^a-z+ ]', '', text)   # + allow kiya
#     return text

# # ---------------- EXACT WORD CHECK ----------------
# def is_exact_word(skill, text):
#     return re.search(rf"\b{re.escape(skill)}\b", text) is not None

# # ---------------- PDF ----------------
# def extract_text_from_pdf(file):
#     reader = PyPDF2.PdfReader(file)
#     text = ""
#     for page in reader.pages:
#         if page.extract_text():
#             text += page.extract_text()
#     return text

# # ---------------- DOCX ----------------
# def extract_text_from_docx(file):
#     doc = Document(file)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + " "
#     return text


# skills_list = [
#     # Programming
#     "python", "java",  "c++", "javascript","php",
#     # AI ENGINEERING
    
#     "deep learning","generative Ai","pytorch","tensorflow","keras","statistics","deep learning framework","CNN", "scikit-learn","Anaconda","Agentic AI"
#     "pandas","numpy","Ai prompting","neural network","Agentic AI",

#     # Data / Analyst
#     "sql", "excel", "power bi", "tableau",
#     "data analysis", "data visualization",
#     "machine learning","matplotlib","basic statistics knowledge",

#     # Web
#     "html", "css", "react", "node js",

#     # HR
#     "hr", "recruitment", "payroll",
#     "employee relations", "talent acquisition",

#     # Marketing
#     "marketing", "seo", "digital marketing",
#     "social media", "content creation", "sales",

#     #cyber security
#     "network security", "penetration testing", "ethical hacking",
    
#     #cloud engineer
#     "aws", "azure", "cloud", "devops",

#     #banking
#     "banking", "finance", "loan", "credit", "risk management",

#     #Accounting
#     "accounting", "tally", "gst", "taxation", "financial analysis",

#     # Finance
#     "accounting", "financial analysis", "budgeting",
    
#     #healthcare
#     "nursing", "patient care", "medical",

#     #graphic designer
#     "photoshop", "illustrator", "design", "creativity",
    
#     #teacher
#     "teaching", "communication", "subject knowledge",

#     # Soft skills
#     "communication", "teamwork", "leadership","digital communication skills",
#     "problem solving","Clear speaking","Problem-solving communication","Verbal Communication","Written Communication",
#     "Listening Skills","Interpersonal Communication","Professional Communication","public speaking skills","persuasion" , "influence skills",
    
#     # Tools
#     "ms office", "word", "powerpoint",
# ]

# def extract_skills(text):
#     text = text.lower()
#     words = text.split()
#     found_skills = []

#     for skill in skills_list:
#         skill_words = skill.split()

#         # ✅ 1. Exact match (best)
#         if f" {skill} " in f" {text} ":
#             found_skills.append(skill)
#             continue

#         # 🔒 2. Short skills → only exact
#         if len(skill) <= 4:
#          if is_exact_word(skill, text):
#           found_skills.append(skill)
#          continue

#         # ✅ 3. Multi-word skills
#         if len(skill_words) > 1:
#             for i in range(len(words) - len(skill_words) + 1):
#                 phrase = " ".join(words[i:i+len(skill_words)])

#                 if fuzz.ratio(skill, phrase) > 90:
#                     found_skills.append(skill)

#         # ✅ 4. Single-word fuzzy (STRICT + SAFE)
#         else:
#             for word in words:

#                 # 🔥 filters
#                 if len(word) < 4:
#                     continue

#                 if abs(len(word) - len(skill)) > 1:
#                     continue

#                 if fuzz.ratio(skill, word) > 92:
                    
#                     # 🔥 EXTRA CHECK (most important)
#                     if word.startswith(skill[:3]):  
#                         found_skills.append(skill)

#     return list(set(found_skills))
# job_skills = {

#     # 💻 Data / IT
#     "Data Analyst": ["python", "sql", "excel", "power bi", "tableau", "data analysis"],
#     "Data Scientist": ["python", "machine learning", "deep learning", "nlp", "statistics"],
#     "Software Developer": ["python", "java", "c++", "javascript", "algorithms"],
#     "Web Developer": ["html", "css", "javascript", "react", "node js"],
#     "AI Engineer": ["machine learning", "deep learning", "tensorflow", "pytorch"],
#     "Cyber Security": ["network security", "penetration testing", "ethical hacking"],
#     "Cloud Engineer": ["aws", "azure", "cloud", "devops"],

#     # 🏢 Business / Management
#     "HR": ["hr", "recruitment", "payroll", "employee relations", "talent acquisition"],
#     "Marketing": ["marketing", "seo", "digital marketing", "sales", "social media"],
#     "Sales": ["sales", "negotiation", "lead generation", "client handling"],
#     "Business Analyst": ["excel", "sql", "business analysis", "requirements gathering"],
#     "Project Manager": ["management", "planning", "leadership", "agile"],

#     # 💰 Finance / Banking
#     "Banking": ["banking", "finance", "loan", "credit", "risk management"],
#     "Accountant": ["accounting", "tally", "gst", "taxation", "financial analysis"],
#     "Financial Analyst": ["finance", "financial modeling", "excel", "budgeting"],

#     # 📞 Support / Operations
#     "Customer Support": ["customer service", "communication", "problem solving"],
#     "Operations": ["operations", "process management", "logistics"],

#     # 🎨 Creative
#     "Graphic Designer": ["photoshop", "illustrator", "design", "creativity"],
#     "Content Writer": ["writing", "content creation", "seo", "blogging"],

#     # 🏥 Others
#     "Teacher": ["teaching", "communication", "subject knowledge"],
#     "Healthcare": ["nursing", "patient care", "medical"],
# }
# if st.button("Analyze Resume"):

#     # 📌 Get text
#     if uploaded_file:
#         if uploaded_file.type == "application/pdf":
#             user_resume = extract_text_from_pdf(uploaded_file)
#         elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             user_resume = extract_text_from_docx(uploaded_file)
#         else:
#             user_resume = str(uploaded_file.read(), "utf-8")

#     # 🧠 Clean + Extract
#     clean_resume = clean_text(user_resume)
#     user_skills = extract_skills(clean_resume)

#     st.subheader("🧠 Skills")
#     for skill in user_skills:
#         st.write("✔", skill)

#     # 🎯 Best Role
#     best_role = ""
#     best_score = 0
#     role_scores = []

#     for role, skills in job_skills.items():
#         matched = set(user_skills) & set(skills)
#         score = len(matched) / len(skills) * 100

#         role_scores.append((role, score))

#         if score > best_score:
#             best_score = score
#             best_role = role

#     st.subheader("🎯 Best Role")
#     st.write(best_role)

#     # 🏆 Top 3
#     role_scores.sort(key=lambda x: x[1], reverse=True)

#     st.subheader("🏆 Top 3 Role Suggestions")
#     for role, score in role_scores[:3]:
#         st.write(f"✔ {role} - {round(score, 2)}%")

#     # 📊 Score
#     st.subheader("📊 Match Score")
#     st.write(str(round(best_score, 2)) + "%")

#     # ❗ Missing Skills
#     if best_role in job_skills:

#         missing = list(set(job_skills[best_role]) - set(user_skills))

#         st.subheader("📌 Missing Skills")

#         if len(missing) == 0:
#             st.success("🎉 No missing skills!")
#         else:
#             for skill in missing:
#                 st.write("•", skill)

#         # 📌 Recommendation
#         st.subheader("📌 Recommendation")

#         if len(missing) == 0:
#             st.success("🎉 Perfect match! You are fully qualified.")
#         elif best_score >= 80:
#             st.info("✔ Good candidate but improvement needed")
#         else:
#             st.warning("⚠ You need improvement")

# import streamlit as st
# import re
# import PyPDF2
# from docx import Document
# from sentence_transformers import SentenceTransformer, util

# # ===============================
# # 🔥 Load AI Model (once only)
# # ===============================
# model = SentenceTransformer('all-MiniLM-L6-v2')

# # ===============================
# # UI
# # ===============================
# st.title("📄 AI Resume Analyzer (ATS Level)")
# st.markdown("### 👨‍💻 Built by: Piyush Solanki")
# st.markdown("GitHub: https://github.com/piyush-tech243")
# st.markdown("LinkedIn: https://www.linkedin.com/in/piyush-solanki-9a5bb2317")
# st.markdown("Made with ❤️ using Streamlit + NLP + AI")

# # ===============================
# # File Upload
# # ===============================
# uploaded_file = st.file_uploader("📤 Upload Resume", type=["pdf", "docx", "txt"])
# user_text = st.text_area("Or paste your resume text")

# # ===============================
# # Extract Text Functions
# # ===============================
# def extract_text_from_pdf(file):
#     reader = PyPDF2.PdfReader(file)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text() or ""
#     return text

# def extract_text_from_docx(file):
#     doc = Document(file)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + " "
#     return text

# # ===============================
# # Clean Text
# # ===============================
# def clean_text(text):
#     text = text.lower()
#     text = re.sub(r'\s+', ' ', text)
#     text = re.sub(r'[^a-z0-9+#. ]', '', text)
#     return text

# # ===============================
# # AI Similarity Function
# # ===============================
# def get_similarity(resume, job_desc):
#     resume_emb = model.encode(resume, convert_to_tensor=True)
#     job_emb = model.encode(job_desc, convert_to_tensor=True)

#     score = util.cos_sim(resume_emb, job_emb)[0][0]
#     return float(score) * 100

# # ===============================
# # JOB ROLES (AI Matching)
# # ===============================
# job_roles = {
#     "Data Analyst": "python sql excel power bi tableau data analysis statistics pandas visualization",
#     "Data Scientist": "python machine learning deep learning nlp tensorflow pytorch statistics sklearn",
#     "Software Developer": "python java c++ javascript data structures algorithms problem solving",
#     "Web Developer": "html css javascript react node js frontend backend web development",
#     "AI Engineer": "machine learning deep learning neural networks pytorch tensorflow ai models",
#     "Cyber Security": "network security ethical hacking penetration testing cybersecurity firewalls",
#     "Cloud Engineer": "aws azure cloud devops docker kubernetes deployment",
#     "Business Analyst": "excel sql business analysis requirements gathering reporting communication",
#     "Graphic Designer": "photoshop illustrator design creativity ui ux branding",
#     "Accountant": "accounting gst tally finance taxation financial reporting",
# }

# # ===============================
# # MAIN LOGIC
# # ===============================
# if st.button("🚀 Analyze Resume"):

#     # ---------------- Extract text ----------------
#     if uploaded_file:
#         if uploaded_file.type == "application/pdf":
#             resume_text = extract_text_from_pdf(uploaded_file)
#         elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             resume_text = extract_text_from_docx(uploaded_file)
#         else:
#             resume_text = str(uploaded_file.read(), "utf-8")
#     else:
#         resume_text = user_text

#     if not resume_text:
#         st.warning("Please upload or paste resume text")
#         st.stop()

#     # ---------------- Clean ----------------
#     clean_resume = clean_text(resume_text)

#     # ---------------- AI Scoring ----------------
#     results = {}

#     for role, desc in job_roles.items():
#         score = get_similarity(clean_resume, desc)
#         results[role] = score

#     # ---------------- Best Role ----------------
#     best_role = max(results, key=results.get)
#     best_score = results[best_role]

#     # ===============================
#     # OUTPUT
#     # ===============================
#     st.subheader("🎯 Best Career Match")
#     st.success(f"{best_role} - {best_score:.2f}% match")

#     st.subheader("📊 ATS Score Progress")
#     st.progress(best_score / 100)

#     # ===============================
#     # Top 3 Roles
#     # ===============================
#     st.subheader("🏆 Top Career Suggestions")

#     sorted_roles = sorted(results.items(), key=lambda x: x[1], reverse=True)

#     for role, score in sorted_roles[:3]:
#         st.write(f"✔ {role} - {score:.2f}%")

#     # ===============================
#     # Insight
#     # ===============================
#     st.subheader("🧠 AI Insight")

#     if best_score > 80:
#         st.success("🔥 Excellent match! You are highly suitable for this role.")
#     elif best_score > 60:
#         st.info("👍 Good match but you can improve skills for better chances.")
#     else:
#         st.warning("⚠ You need skill improvement for better job matching.")

import streamlit as st
import re
import PyPDF2
from docx import Document
from rapidfuzz import fuzz

st.title("📄 Resume analyzer.py")
st.markdown("### 👨‍💻 Built by: Piyush Solanki")
st.markdown("GitHub: https://github.com/piyush-tech243")
st.markdown("🔗 LinkedIn: https://www.linkedin.com/in/piyush-solanki-9a5bb2317")
st.markdown("Made with ❤️ using Python, Streamlit, RapidFuzz, PyPDF2 & NLP")

# ---------------- FILE INPUT ----------------
uploaded_file = st.file_uploader("📤 Upload Resume (Max 5MB)", type=["pdf", "docx", "txt"])
MAX_FILE_SIZE = 5 * 1024 * 1024   # 5MB
user_resume = st.text_area("Or paste your Text here")

# ---------------- CLEAN TEXT ----------------
def clean_text(text):
    text = text.lower()
    text = re.sub(r'[-_/]', ' ', text)
    text = re.sub(r'[()]', ' ', text)
    text = re.sub(r'[^a-z0-9+ ]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# ---------------- NORMALIZE ----------------
def normalize_word(word):
    word = word.lower().strip()
    if word.endswith("s") and len(word) > 3:
        word = word[:-1]
    return word

# ---------------- FILE READ ----------------
def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text()
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + " "
    return text

# ---------------- SKILLS ----------------
skills_list = list(set([
     "python", "java",  "c++", "javascript","php",
    # AI ENGINEERING
    
    "deep learning","generative Ai","pytorch","tensorflow","keras","statistics","deep learning framework","CNN", "scikit-learn","Anaconda","Agentic AI",
    "pandas","numpy","Ai prompting","neural network",

    # Data / Analyst
    "sql", "excel", "power bi", "tableau",
    "data analysis", "data visualization",
    "machine learning","matplotlib","basic statistics knowledge",
     "Data Cleaning", "Exploratory Data Analysis (EDA)", "Insight Generation", "Descriptive Statistics",
    # Web
    "html", "css", "react", "node js",

    # HR
    "hr", "recruitment", "payroll",
    "employee relations", "talent acquisition",

    # Marketing
    "marketing", "seo", "digital marketing",
    "social media", "content creation", "sales",

    #all skills
    "machine learning","deep learning","generative ai","artificial intelligence","ai",
    "pytorch","tensorflow","keras","cnn","neural network",
    "scikit learn","scikit-learn","numpy","pandas","matplotlib",
    "anaconda","agentic ai",
    "sql","excel","power bi","tableau",
    "data analysis","data visualization","data cleaning",
    "exploratory data analysis","eda","insight generation",
    "statistics","descriptive statistics",
    "html","css","react","node js",
    "aws","azure","cloud","devops",
    "communication","communication skills","teamwork","leadership","problem solving"

    #cyber security
    "network security", "penetration testing", "ethical hacking",
    
    #cloud engineer
    "aws", "azure", "cloud", "devops",

    #banking
    "banking", "finance", "loan", "credit", "risk management",

    #Accounting
    "accounting", "tally", "gst", "taxation", "financial analysis",

    # Finance
    "accounting", "financial analysis", "budgeting",
    
    #healthcare
    "nursing", "patient care", "medical",

    #graphic designer
    "photoshop", "illustrator", "design", "creativity",
    
    #teacher
    "teaching", "communication", "subject knowledge",

    # Soft skills
     "teamwork", "leadership","digital communication skills",
    "problem solving","Clear speaking","Problem-solving communication","Verbal Communication","Written Communication",
    "Listening Skills","Interpersonal Communication","Professional Communication","public speaking skills","persuasion" , "influence skills",
    
    # Tools
    "ms office", "word", "powerpoint",
]))

# ---------------- SKILL EXTRACTION ----------------
def extract_skills(text):
    words = text.split()
    found_skills = []

    for skill in skills_list:
        skill_words = skill.split()

        # Exact match
        if f" {skill} " in f" {text} ":
            found_skills.append(skill)
            continue

        # Short skills (AI, CNN)
        if len(skill) <= 3:
            if f" {skill} " in f" {text} ":
                found_skills.append(skill)
            continue

        # Multi-word skills
        if len(skill_words) > 1:
            for i in range(len(words) - len(skill_words) + 1):
                phrase = " ".join(words[i:i+len(skill_words)])
                if fuzz.token_sort_ratio(skill, phrase) > 85:
                    found_skills.append(skill)

        # Single-word skills
        else:
            for word in words:
                word = normalize_word(word)

                if len(word) < 3:
                    continue

                # 🔥 strict check (false positive control)
                if skill not in text:
                    continue

                if fuzz.partial_ratio(skill, word) > 92:
                    if word.startswith(skill[:3]):
                        found_skills.append(skill)

    return list(set(found_skills))

# ---------------- FINAL CLEANING ----------------
def clean_final_skills(skills):
    normalized = []

    for skill in skills:
        skill = skill.lower().strip()

        if skill in ["eda", "exploratory data analysis"]:
            skill = "data analysis"

        elif skill in ["descriptive statistics"]:
            skill = "statistics"

        elif skill in ["communication skills"]:
            skill = "communication"

        elif skill in ["scikit learn"]:
            skill = "scikit-learn"

        elif skill == "ai":
            skill = "artificial intelligence"

        normalized.append(skill)

    return sorted(list(set(normalized)))

# ---------------- JOB ROLES ----------------
job_skills = {
    "Data Analyst": ["python","sql","excel","power bi","tableau","data analysis"],
    "Data Scientist": ["python","machine learning","deep learning","statistics"],
    "AI Engineer": ["machine learning","deep learning","tensorflow","pytorch"],
    "Software Developer": ["python","java","c++","javascript"],
    "Web Developer": ["html","css","javascript","react","node js"],
    "Cloud Engineer": ["aws","azure","cloud","devops"]
}

# ---------------- MAIN ----------------
if st.button("Analyze Resume"):

    if uploaded_file:
        if uploaded_file.size > MAX_FILE_SIZE:
         st.error("❌ File size exceeds 5MB limit. Please upload a file under 5MB.")
         st.stop()
        if uploaded_file.type == "application/pdf":
            user_resume = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            user_resume = extract_text_from_docx(uploaded_file)
        else:
            user_resume = str(uploaded_file.read(), "utf-8")

    clean_resume = clean_text(user_resume)

    user_skills = extract_skills(clean_resume)
    user_skills = clean_final_skills(user_skills)

    # Skills
    st.subheader("🧠 Skills Detected")
    for skill in user_skills:
        st.write("✔", skill)

    # Best Role
    best_role = ""
    best_score = 0
    role_scores = []

    for role, skills in job_skills.items():
        matched = set(user_skills) & set(skills)
        score = len(matched) / len(skills) * 100
        role_scores.append((role, score))

        if score > best_score:
            best_score = score
            best_role = role

    st.subheader("🎯 Best Role")
    st.write(best_role)

    # Top 3
    role_scores.sort(key=lambda x: x[1], reverse=True)

    st.subheader("🏆 Top 3 Role Suggestions")
    for role, score in role_scores[:3]:
        st.write(f"✔ {role} - {round(score,2)}%")

    # Score
    st.subheader("📊 Match Score")
    st.write(f"{round(best_score,2)}%")

    # Missing Skills
    if best_role in job_skills:
        missing = list(set(job_skills[best_role]) - set(user_skills))

        st.subheader("📌 Missing Skills")
        if not missing:
            st.success("🎉 No missing skills!")
        else:
            for skill in missing:
                st.write("•", skill)

        # Recommendation
        st.subheader("📌 Recommendation")
        if not missing:
            st.success("🎉 Perfect match!")
        elif best_score >= 80:
            st.info("✔ Good candidate but improvement needed")
        else:
            st.warning("⚠ You need improvement")